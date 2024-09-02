import streamlit as st
import whisperx
import ffmpeg
import io
import subprocess
from datetime import datetime
from docx import Document
import torch
import account

def log_activity(message):
    """Logs activity with a timestamp to Streamlit's console."""
    timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    st.write(f"{timestamp} - {message}")

def format_duration(seconds):
    """Converts seconds to a human-readable format like '1 min 29 sec'."""
    minutes, seconds = divmod(seconds, 60)
    if minutes > 0:
        return f"{int(minutes)} min {int(seconds)} sec"
    else:
        return f"{int(seconds)} sec"

def format_timestamp(seconds):
    """Converts seconds to HH:MM:SS format."""
    hours, remainder = divmod(int(seconds), 3600)
    minutes, seconds = divmod(remainder, 60)
    return f"{hours:02}:{minutes:02}:{seconds:02}"

def extract_audio(file_buffer):
    """Extracts audio from a video file and returns it as an in-memory .wav file."""
    output_audio = io.BytesIO()
    log_activity(f"Calling extract_audio for {file_buffer.name}.")

    try:
        # Use 'pipe:' for both input and output to handle in-memory files
        process = (
            ffmpeg
            .input('pipe:', format=file_buffer.name.split('.')[-1])  # Detect format from filename
            .output('pipe:', format='wav')
            .run_async(pipe_stdin=True, pipe_stdout=True, pipe_stderr=True)
        )
        # Read from file_buffer and write to process.stdin
        stdout, stderr = process.communicate(input=file_buffer.read())
        output_audio.write(stdout)
        output_audio.seek(0)
        log_activity("Audio extraction complete.")
        return output_audio
    except Exception as e:
        log_activity(f"Error extracting audio: {e}")
        return None

def apply_noise_reduction(input_audio):
    """Applies noise reduction using SoX and returns the cleaned audio as an in-memory buffer."""
    cleaned_audio = io.BytesIO()
    log_activity("Applying noise reduction.")

    try:
        process = subprocess.Popen(
            ['sox', '-t', 'wav', '-', '-t', 'wav', '-', 'noisered', '0.15'],
            stdin=subprocess.PIPE,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE
        )
        stdout, stderr = process.communicate(input=input_audio.read())
        if process.returncode == 0:
            cleaned_audio.write(stdout)
            cleaned_audio.seek(0)
            log_activity("Noise reduction applied.")
            return cleaned_audio
        else:
            log_activity(f"Error during noise reduction: {stderr.decode()}")
            return None
    except subprocess.CalledProcessError as e:
        log_activity(f"Error during noise reduction: {e}")
        return None

def transcribe_audio_with_whisperx(audio_file):
    """Transcribes the given audio file using WhisperX and returns the transcription as an in-memory .docx file."""
    log_activity("Starting transcription.")

    device = "cuda" if torch.cuda.is_available() else "cpu"
    batch_size = 8
    compute_type = "float16" if torch.cuda.is_available() else "float32"

    try:
        model = whisperx.load_model("large-v2", device, compute_type=compute_type, language='en')
        log_activity("Whisper model loaded.")

        audio = whisperx.load_audio(audio_file)
        result = model.transcribe(audio, batch_size=batch_size)
        log_activity("Transcription completed.")

        model_a, metadata = whisperx.load_align_model(language_code=result["language"], device=device)
        aligned_result = whisperx.align(result["segments"], model_a, metadata, audio, device, return_char_alignments=False)
        log_activity("Alignment completed.")

        diarize_model = whisperx.DiarizationPipeline(use_auth_token="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", device=device)
        diarize_segments = diarize_model(audio)
        final_result = whisperx.assign_word_speakers(diarize_segments, aligned_result)
        log_activity("Speaker diarization completed.")

        docx_buffer = io.BytesIO()
        document = Document()

        current_speaker = None
        current_segment_start = None
        current_segment_end = None
        segment_text = []

        for segment in final_result["segments"]:
            start = segment["start"]
            end = segment["end"]
            speaker = segment.get("speaker", "Unknown Speaker")
            text = segment["text"]

            if speaker != current_speaker:
                if current_speaker is not None:
                    document.add_paragraph(f"[{format_timestamp(current_segment_start)} - {format_timestamp(current_segment_end)}] {current_speaker}: {' '.join(segment_text)}\n")
                current_speaker = speaker
                current_segment_start = start
                current_segment_end = end
                segment_text = [text]
            else:
                current_segment_end = end
                segment_text.append(text)

        if current_speaker is not None:
            document.add_paragraph(f"[{format_timestamp(current_segment_start)} - {format_timestamp(current_segment_end)}] {current_speaker}: {' '.join(segment_text)}\n")

        document.save(docx_buffer)
        docx_buffer.seek(0)
        log_activity("Transcription saved.")
        return docx_buffer

    except Exception as e:
        log_activity(f"Error during transcription: {e}")
        return None

def app():
    st.title("Transcription Service")
    st.write("Upload an audio or video file, and get a transcribed .docx file with speaker diarization.")

    uploaded_file = st.file_uploader("Choose an audio or video file", type=["wav", "mp4", "mkv", "avi", "mp3"])

    if uploaded_file is not None:
        st.success(f"File {uploaded_file.name} uploaded successfully!")

        # Add a button to start the transcription process
        if st.button("Transcribe"):
            st.write("Processing the file...")
            audio_data = extract_audio(uploaded_file)
            if audio_data:
                audio_input = apply_noise_reduction(audio_data)
                if audio_input:
                    st.write("Transcribing the audio...")
                    docx_file = transcribe_audio_with_whisperx(audio_input)

                    if docx_file:
                        st.write("Transcription complete!")
                        st.download_button(
                            label="Download Transcription",
                            data=docx_file,
                            file_name=f"{uploaded_file.name.split('.')[0]}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    else:
                        st.error("Failed to transcribe the audio.")
                else:
                    st.error("Failed to apply noise reduction.")
            else:
                st.error("Failed to extract audio.")

app()
